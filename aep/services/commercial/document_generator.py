from pathlib import Path
from docx import Document

from .context import CommercialOfferContext


class DocumentGenerator:
    def __init__(self, output_dir="output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)

    def generate(self, context: CommercialOfferContext) -> Path:
        template_path = Path(context.template_path)

        if not template_path.exists():
            raise FileNotFoundError(f"Шаблон не найден: {template_path}")

        doc = Document(template_path)

        self._append_test_block(doc, context)

        file_path = self.output_dir / "commercial_offer_from_template.docx"
        doc.save(file_path)

        return file_path

    def _append_test_block(self, doc: Document, context: CommercialOfferContext) -> None:
        doc.add_paragraph("")
        doc.add_paragraph("=== AEP TEST DATA ===")
        doc.add_paragraph(f"Компания: {context.request.company}")
        doc.add_paragraph(f"Заказчик: {context.request.customer}")
        doc.add_paragraph(f"Продукт: {context.request.product_name}")
        doc.add_paragraph(f"Марка угля: {context.request.coal_mark}")
        doc.add_paragraph(f"Объем: {context.request.volume_tons} тонн")
        doc.add_paragraph(f"Пункт поставки: {context.request.destination}")
        doc.add_paragraph(f"Цена без НДС: {context.calculated_price_without_vat}")
        doc.add_paragraph(f"Цена с НДС: {context.calculated_price_with_vat}")
        doc.add_paragraph(f"Сертификат: {context.certificate_path}")